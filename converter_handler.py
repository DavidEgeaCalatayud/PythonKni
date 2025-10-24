from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from PIL import Image
from docx import Document
import os
import fitz  # PyMuPDF (para PDF → imágenes)

# ---------------- IMÁGENES ----------------
def images_to_pdf(image_files, output_file):
    """Convierte una lista de imágenes en un PDF."""
    c = canvas.Canvas(output_file, pagesize=A4)
    width, height = A4

    for img_path in image_files:
        try:
            img = Image.open(img_path)
            img_width, img_height = img.size
            aspect = img_height / img_width
            new_width = width
            new_height = width * aspect
            if new_height > height:
                new_height = height
                new_width = height / aspect
            c.drawImage(img_path, 0, height - new_height, new_width, new_height)
            c.showPage()
        except Exception as e:
            print(f"Error con {img_path}: {e}")
    c.save()

def pdf_to_images(pdf_file, output_folder):
    """Convierte un PDF en imágenes (PNG)."""
    doc = fitz.open(pdf_file)
    saved_files = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap()
        output_path = os.path.join(output_folder, f"page_{i+1}.png")
        pix.save(output_path)
        saved_files.append(output_path)
    return saved_files

# ---------------- TEXTO Y WORD ----------------
def text_to_docx(text_file, output_file):
    """Convierte un archivo TXT en DOCX."""
    doc = Document()
    with open(text_file, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            doc.add_paragraph(line.strip())
    doc.save(output_file)

def docx_to_text(docx_file, output_file):
    """Convierte un archivo DOCX en TXT."""
    doc = Document(docx_file)
    with open(output_file, "w", encoding="utf-8") as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")

def docx_to_pdf(docx_file, output_file):
    """Convierte un DOCX en PDF (simplificado como texto plano)."""
    doc = Document(docx_file)
    c = canvas.Canvas(output_file, pagesize=A4)
    width, height = A4
    y = height - 50

    for para in doc.paragraphs:
        text = para.text
        c.drawString(50, y, text)
        y -= 15
        if y < 50:
            c.showPage()
            y = height - 50
    c.save()

from PyQt5.QtWidgets import (
    QMainWindow, QVBoxLayout, QWidget, QPushButton,
    QFileDialog, QMessageBox
)
from converter_handler import (
    images_to_pdf, pdf_to_images, text_to_docx,
    docx_to_text, docx_to_pdf
)
import os

class ConverterWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Convertidor de Archivos")
        self.setGeometry(200, 200, 400, 300)

        layout = QVBoxLayout()

        # Opciones de conversión
        btn_img_pdf = QPushButton("Imágenes → PDF")
        btn_img_pdf.clicked.connect(self.convert_images_to_pdf)
        layout.addWidget(btn_img_pdf)

        btn_pdf_img = QPushButton("PDF → Imágenes")
        btn_pdf_img.clicked.connect(self.convert_pdf_to_images)
        layout.addWidget(btn_pdf_img)

        btn_txt_docx = QPushButton("TXT → DOCX")
        btn_txt_docx.clicked.connect(self.convert_text_to_docx)
        layout.addWidget(btn_txt_docx)

        btn_docx_txt = QPushButton("DOCX → TXT")
        btn_docx_txt.clicked.connect(self.convert_docx_to_text)
        layout.addWidget(btn_docx_txt)

        btn_docx_pdf = QPushButton("DOCX → PDF")
        btn_docx_pdf.clicked.connect(self.convert_docx_to_pdf)
        layout.addWidget(btn_docx_pdf)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

    # ---------------- FUNCIONES ----------------
    def convert_images_to_pdf(self):
        files, _ = QFileDialog.getOpenFileNames(self, "Seleccionar imágenes", "", "Imágenes (*.png *.jpg *.jpeg *.bmp)")
        if not files:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Guardar PDF", "", "PDF (*.pdf)")
        if not save_path:
            return
        images_to_pdf(files, save_path)
        QMessageBox.information(self, "Conversión completada", f"PDF creado en:\n{save_path}")

    def convert_pdf_to_images(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Seleccionar PDF", "", "PDF (*.pdf)")
        if not file_path:
            return
        folder = QFileDialog.getExistingDirectory(self, "Seleccionar carpeta de destino")
        if not folder:
            return
        saved_files = pdf_to_images(file_path, folder)
        QMessageBox.information(self, "Conversión completada", f"Se han guardado {len(saved_files)} imágenes en:\n{folder}")

    def convert_text_to_docx(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Seleccionar TXT", "", "Texto (*.txt)")
        if not file_path:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Guardar DOCX", "", "Word (*.docx)")
        if not save_path:
            return
        text_to_docx(file_path, save_path)
        QMessageBox.information(self, "Conversión completada", f"DOCX creado en:\n{save_path}")

    def convert_docx_to_text(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Seleccionar DOCX", "", "Word (*.docx)")
        if not file_path:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Guardar TXT", "", "Texto (*.txt)")
        if not save_path:
            return
        docx_to_text(file_path, save_path)
        QMessageBox.information(self, "Conversión completada", f"TXT creado en:\n{save_path}")

    def convert_docx_to_pdf(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Seleccionar DOCX", "", "Word (*.docx)")
        if not file_path:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Guardar PDF", "", "PDF (*.pdf)")
        if not save_path:
            return
        docx_to_pdf(file_path, save_path)
        QMessageBox.information(self, "Conversión completada", f"PDF creado en:\n{save_path}")
