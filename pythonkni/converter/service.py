from __future__ import annotations
from .models import (
    ConversionResult,
)
import os
import shutil
import tempfile
from dataclasses import dataclass
from pathlib import Path
import logging
from xml.dom.minidom import Document as XMLDocument
import fitz
from docx import Document
from PIL import Image
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from pythonkni.core.tasks import WorkerCancelled


class OutputTransaction:
    """Stage one or more outputs and publish them atomically as a logical batch.

    Staging and backups live in the destination directory so ``os.replace`` stays
    on the same filesystem. Existing outputs are only moved aside during commit;
    if any publish step fails, every previously published output is rolled back.
    """

    def __init__(self, destination_dir: str | Path):
        self.destination_dir = Path(destination_dir)
        self.destination_dir.mkdir(parents=True, exist_ok=True)
        self.staging_dir = Path(
            tempfile.mkdtemp(prefix=".pythonkni-converter-", dir=self.destination_dir)
        )
        self._entries: list[tuple[Path, Path]] = []
        self._committed = False

    def stage_for(self, final_path: str | Path) -> Path:
        final = Path(final_path)
        if final.parent.resolve(strict=False) != self.destination_dir.resolve(strict=False):
            raise ValueError("Todos los destinos de una transacción deben compartir carpeta.")
        index = len(self._entries)
        stage = self.staging_dir / f"{index:04d}_{final.stem}.stage{final.suffix}"
        self._entries.append((stage, final))
        return stage

    def commit(self) -> list[str]:
        states: list[tuple[Path, Path | None]] = []
        try:
            for index, (stage, final) in enumerate(self._entries):
                if not stage.exists():
                    raise FileNotFoundError(f"No existe el resultado temporal: {stage}")

                backup = None
                if final.exists():
                    backup = self.staging_dir / f"{index:04d}_{final.name}.backup"
                    os.replace(final, backup)
                states.append((final, backup))
                os.replace(stage, final)
        except Exception:
            for final, backup in reversed(states):
                try:
                    if final.exists():
                        final.unlink()
                    if backup is not None and backup.exists():
                        os.replace(backup, final)
                except OSError:
                    pass
            raise

        self._committed = True
        outputs = [str(final) for _, final in self._entries]
        self._cleanup()
        return outputs

    def _cleanup(self) -> None:
        shutil.rmtree(self.staging_dir, ignore_errors=True)

    def abort(self) -> None:
        if not self._committed:
            self._cleanup()

    def __enter__(self) -> "OutputTransaction":
        return self

    def __exit__(self, exc_type, exc_value, traceback) -> None:
        if not self._committed:
            self.abort()


logger = logging.getLogger(__name__)


def validate_extension(file_path: str, allowed_extensions: set[str]) -> bool:
    return os.path.splitext(file_path)[1].lower() in allowed_extensions


def _check_worker(worker):
    if worker is not None:
        worker.check_cancelled()


def _report(worker, message, current=None, total=None):
    if worker is None:
        return
    payload = {"message": message}
    if current is not None and total:
        payload["percent"] = int((current / total) * 100)
    worker.report_progress(payload)


def _single_output_transaction(output_file: str | os.PathLike[str]) -> OutputTransaction:
    output = Path(output_file)
    return OutputTransaction(output.parent)


def images_to_pdf(image_files, output_file, worker=None):
    """Convierte imágenes a PDF sin publicar resultados parciales."""
    output = Path(output_file)
    warnings = []
    added_pages = 0

    with _single_output_transaction(output) as transaction:
        stage = transaction.stage_for(output)
        pdf_canvas = canvas.Canvas(str(stage), pagesize=A4)
        width, height = A4
        total = len(image_files)

        try:
            for index, img_path in enumerate(image_files, start=1):
                _check_worker(worker)
                try:
                    with Image.open(img_path) as image:
                        img_width, img_height = image.size
                        if img_width <= 0 or img_height <= 0:
                            raise ValueError("Dimensiones de imagen no válidas.")

                    aspect = img_height / img_width
                    new_width = width
                    new_height = width * aspect
                    if new_height > height:
                        new_height = height
                        new_width = height / aspect
                    pdf_canvas.drawImage(
                        img_path,
                        0,
                        height - new_height,
                        new_width,
                        new_height,
                    )
                    pdf_canvas.showPage()
                    added_pages += 1
                except Exception as error:
                    warning = f"{img_path}: {error}"
                    warnings.append(warning)
                    logger.warning("No se pudo añadir la imagen %s al PDF", img_path, exc_info=True)
                _report(worker, f"Añadiendo imagen {index}/{total}", index, total)

            _check_worker(worker)
            pdf_canvas.save()
            _check_worker(worker)
        except Exception:
            raise

        if added_pages == 0:
            return ConversionResult.failed(
                "No se pudo añadir ninguna imagen al PDF; no se publicó ningún archivo.",
                warnings=warnings,
            )

        outputs = transaction.commit()
        return ConversionResult.completed(outputs, warnings=warnings)


def pdf_to_images(pdf_file, output_folder, worker=None):
    """Convierte un PDF en PNG como una única transacción de lote."""
    output_dir = Path(output_folder)
    document = fitz.open(pdf_file)
    try:
        total = len(document)
        if total == 0:
            return ConversionResult.failed("El PDF no contiene páginas para convertir.")

        with OutputTransaction(output_dir) as transaction:
            for index, page in enumerate(document, start=1):
                _check_worker(worker)
                final_path = output_dir / f"page_{index}.png"
                stage = transaction.stage_for(final_path)
                pix = page.get_pixmap()
                pix.save(str(stage))
                _report(worker, f"Renderizando página {index}/{total}", index, total)

            _check_worker(worker)
            outputs = transaction.commit()
            return ConversionResult.completed(outputs)
    finally:
        document.close()


def text_to_docx(text_file, output_file, worker=None):
    """Convierte un TXT en DOCX mediante staging + publicación atómica."""
    output = Path(output_file)
    document = Document()
    with open(text_file, "r", encoding="utf-8", errors="ignore") as file:
        for index, line in enumerate(file, start=1):
            _check_worker(worker)
            document.add_paragraph(line.strip())
            if index % 100 == 0:
                _report(worker, f"Procesando línea {index}")

    _check_worker(worker)
    with _single_output_transaction(output) as transaction:
        stage = transaction.stage_for(output)
        document.save(str(stage))
        _check_worker(worker)
        return ConversionResult.completed(transaction.commit())


def docx_to_text(docx_file, output_file, worker=None):
    """Convierte DOCX a TXT sin tocar el destino hasta terminar."""
    output = Path(output_file)
    document = Document(docx_file)
    total = len(document.paragraphs)

    with _single_output_transaction(output) as transaction:
        stage = transaction.stage_for(output)
        with stage.open("w", encoding="utf-8") as file:
            for index, paragraph in enumerate(document.paragraphs, start=1):
                _check_worker(worker)
                file.write(paragraph.text + "\n")
                _report(worker, f"Procesando párrafo {index}/{total}", index, total)
            file.flush()
            os.fsync(file.fileno())

        _check_worker(worker)
        return ConversionResult.completed(transaction.commit())


def docx_to_pdf(docx_file, output_file, worker=None):
    """Convierte DOCX en PDF simplificado con publicación atómica."""
    output = Path(output_file)
    document = Document(docx_file)
    total = len(document.paragraphs)

    with _single_output_transaction(output) as transaction:
        stage = transaction.stage_for(output)
        pdf_canvas = canvas.Canvas(str(stage), pagesize=A4)
        _, height = A4
        y = height - 50

        for index, paragraph in enumerate(document.paragraphs, start=1):
            _check_worker(worker)
            pdf_canvas.drawString(50, y, paragraph.text)
            y -= 15
            if y < 50:
                pdf_canvas.showPage()
                y = height - 50
            _report(worker, f"Procesando párrafo {index}/{total}", index, total)

        _check_worker(worker)
        pdf_canvas.save()
        _check_worker(worker)
        return ConversionResult.completed(transaction.commit())


def text_to_kml(txt_file, output_file, worker=None):
    """Convierte TXT (lat,lon,nombre opcional) a KML de forma atómica."""
    output = Path(output_file)
    document = XMLDocument()
    kml = document.createElement("kml")
    kml.setAttribute("xmlns", "http://www.opengis.net/kml/2.2")
    document.appendChild(kml)
    root = document.createElement("Document")
    kml.appendChild(root)

    with open(txt_file, "r", encoding="utf-8", errors="ignore") as file:
        for index, line in enumerate(file, start=1):
            _check_worker(worker)
            parts = line.strip().replace("\t", ",").replace(";", ",").split(",")
            if len(parts) >= 2:
                lat, lon = parts[0].strip(), parts[1].strip()
                name = parts[2].strip() if len(parts) > 2 else f"Punto ({lat},{lon})"
                placemark = document.createElement("Placemark")
                pname = document.createElement("name")
                pname.appendChild(document.createTextNode(name))
                placemark.appendChild(pname)
                point = document.createElement("Point")
                coords = document.createElement("coordinates")
                coords.appendChild(document.createTextNode(f"{lon},{lat},0"))
                point.appendChild(coords)
                placemark.appendChild(point)
                root.appendChild(placemark)
            if index % 100 == 0:
                _report(worker, f"Procesando línea {index}")

    _check_worker(worker)
    with _single_output_transaction(output) as transaction:
        stage = transaction.stage_for(output)
        with stage.open("w", encoding="utf-8") as file:
            file.write(document.toprettyxml(indent="  "))
            file.flush()
            os.fsync(file.fileno())
        _check_worker(worker)
        return ConversionResult.completed(transaction.commit())


def kml_to_text(kml_file, output_file, worker=None):
    """Convierte KML a TXT (lat,lon,nombre) mediante staging."""
    from xml.dom import minidom

    output = Path(output_file)
    document = minidom.parse(kml_file)
    placemarks = document.getElementsByTagName("Placemark")
    total = len(placemarks)

    with _single_output_transaction(output) as transaction:
        stage = transaction.stage_for(output)
        with stage.open("w", encoding="utf-8") as file:
            for index, placemark in enumerate(placemarks, start=1):
                _check_worker(worker)
                names = placemark.getElementsByTagName("name")
                name = names[0].firstChild.nodeValue if names else ""
                coords_nodes = placemark.getElementsByTagName("coordinates")
                if not coords_nodes or coords_nodes[0].firstChild is None:
                    raise ValueError(f"Placemark {index} sin coordenadas.")
                coords = coords_nodes[0].firstChild.nodeValue.strip()
                lon, lat, *_ = coords.split(",")
                file.write(f"{lat},{lon},{name}\n")
                _report(worker, f"Procesando punto {index}/{total}", index, total)
            file.flush()
            os.fsync(file.fileno())

        _check_worker(worker)
        return ConversionResult.completed(transaction.commit())


def conversion_task(worker, function, args):
    """Run one conversion and turn ordinary failures into structured results."""
    try:
        worker.check_cancelled()
        result = function(*args, worker=worker)
        if not isinstance(result, ConversionResult):
            result = ConversionResult.completed([str(result)])
        return result
    except WorkerCancelled:
        raise
    except Exception as error:
        logger.exception("Conversion failed")
        return ConversionResult.failed(str(error))


def batch_conversion_task(worker, function, jobs):
    """Convert a folder as all-or-nothing output publication."""
    if not jobs:
        return ConversionResult.failed("No hay archivos para convertir.")

    final_paths = [Path(args[-1]) for args in jobs]
    destination_dir = final_paths[0].parent
    if any(
        path.parent.resolve(strict=False) != destination_dir.resolve(strict=False)
        for path in final_paths
    ):
        return ConversionResult.failed(
            "Todos los destinos del lote deben estar en la misma carpeta."
        )

    warnings = []
    total = len(jobs)
    with OutputTransaction(destination_dir) as transaction:
        try:
            for index, (args, final_path) in enumerate(zip(jobs, final_paths), start=1):
                worker.check_cancelled()
                stage = transaction.stage_for(final_path)
                staged_args = (*args[:-1], str(stage))
                result = function(*staged_args, worker=worker)
                if not isinstance(result, ConversionResult):
                    result = ConversionResult.completed([str(result)])
                warnings.extend(result.warnings)
                if not result.success:
                    failures = result.failures or (f"Falló la conversión de {args[0]}",)
                    return ConversionResult.failed(*failures, warnings=warnings)
                worker.report_progress(
                    {
                        "message": f"Archivo {index}/{total}",
                        "percent": int((index / total) * 100),
                    }
                )

            worker.check_cancelled()
            return ConversionResult.completed(transaction.commit(), warnings=warnings)
        except WorkerCancelled:
            raise
        except Exception as error:
            logger.exception("Batch conversion failed")
            return ConversionResult.failed(str(error), warnings=warnings)
