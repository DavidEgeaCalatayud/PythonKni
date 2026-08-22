from __future__ import annotations

import logging
import os
from pathlib import Path
from xml.dom.minidom import Document as XMLDocument

import fitz
from docx import Document
from PIL import Image
from PyQt5.QtCore import QTimer
from PyQt5.QtWidgets import (
    QFileDialog,
    QHBoxLayout,
    QLabel,
    QMessageBox,
    QPushButton,
    QVBoxLayout,
    QWidget,
)
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from pythonkni.core.tasks import WorkerCancelled
from tools.base_tool import BaseTool
from tools.converter_outputs import ConversionResult, OutputTransaction
from tools.worker import Worker


logger = logging.getLogger(__name__)


def validate_extension(file_path: str, allowed_extensions: set[str]) -> bool:
    return os.path.splitext(file_path)[1].lower() in allowed_extensions


def _check_worker(worker):
    if worker is not None:
        worker.check_cancelled()


def _report(worker, message, current=None, total=None):
    if worker is None:
        return
    payload = {"message": message}
    if current is not None and total:
        payload["percent"] = int((current / total) * 100)
    worker.report_progress(payload)


def _single_output_transaction(output_file: str | os.PathLike[str]) -> OutputTransaction:
    output = Path(output_file)
    return OutputTransaction(output.parent)


# ---------------- IMÁGENES ----------------
def images_to_pdf(image_files, output_file, worker=None):
    """Convierte imágenes a PDF sin publicar resultados parciales."""
    output = Path(output_file)
    warnings = []
    added_pages = 0

    with _single_output_transaction(output) as transaction:
        stage = transaction.stage_for(output)
        pdf_canvas = canvas.Canvas(str(stage), pagesize=A4)
        width, height = A4
        total = len(image_files)

        try:
            for index, img_path in enumerate(image_files, start=1):
                _check_worker(worker)
                try:
                    with Image.open(img_path) as image:
                        img_width, img_height = image.size
                        if img_width <= 0 or img_height <= 0:
                            raise ValueError("Dimensiones de imagen no válidas.")

                    aspect = img_height / img_width
                    new_width = width
                    new_height = width * aspect
                    if new_height > height:
                        new_height = height
                        new_width = height / aspect
                    pdf_canvas.drawImage(
                        img_path,
                        0,
                        height - new_height,
                        new_width,
                        new_height,
                    )
                    pdf_canvas.showPage()
                    added_pages += 1
                except Exception as error:
                    warning = f"{img_path}: {error}"
                    warnings.append(warning)
                    logger.warning("No se pudo añadir la imagen %s al PDF", img_path, exc_info=True)
                _report(worker, f"Añadiendo imagen {index}/{total}", index, total)

            _check_worker(worker)
            pdf_canvas.save()
            _check_worker(worker)
        except Exception:
            raise

        if added_pages == 0:
            return ConversionResult.failed(
                "No se pudo añadir ninguna imagen al PDF; no se publicó ningún archivo.",
                warnings=warnings,
            )

        outputs = transaction.commit()
        return ConversionResult.completed(outputs, warnings=warnings)


def pdf_to_images(pdf_file, output_folder, worker=None):
    """Convierte un PDF en PNG como una única transacción de lote."""
    output_dir = Path(output_folder)
    document = fitz.open(pdf_file)
    try:
        total = len(document)
        if total == 0:
            return ConversionResult.failed("El PDF no contiene páginas para convertir.")

        with OutputTransaction(output_dir) as transaction:
            for index, page in enumerate(document, start=1):
                _check_worker(worker)
                final_path = output_dir / f"page_{index}.png"
                stage = transaction.stage_for(final_path)
                pix = page.get_pixmap()
                pix.save(str(stage))
                _report(worker, f"Renderizando página {index}/{total}", index, total)

            _check_worker(worker)
            outputs = transaction.commit()
            return ConversionResult.completed(outputs)
    finally:
        document.close()


# ---------------- TEXTO Y WORD ----------------
def text_to_docx(text_file, output_file, worker=None):
    """Convierte un TXT en DOCX mediante staging + publicación atómica."""
    output = Path(output_file)
    document = Document()
    with open(text_file, "r", encoding="utf-8", errors="ignore") as file:
        for index, line in enumerate(file, start=1):
            _check_worker(worker)
            document.add_paragraph(line.strip())
            if index % 100 == 0:
                _report(worker, f"Procesando línea {index}")

    _check_worker(worker)
    with _single_output_transaction(output) as transaction:
        stage = transaction.stage_for(output)
        document.save(str(stage))
        _check_worker(worker)
        return ConversionResult.completed(transaction.commit())


def docx_to_text(docx_file, output_file, worker=None):
    """Convierte DOCX a TXT sin tocar el destino hasta terminar."""
    output = Path(output_file)
    document = Document(docx_file)
    total = len(document.paragraphs)

    with _single_output_transaction(output) as transaction:
        stage = transaction.stage_for(output)
        with stage.open("w", encoding="utf-8") as file:
            for index, paragraph in enumerate(document.paragraphs, start=1):
                _check_worker(worker)
                file.write(paragraph.text + "\n")
                _report(worker, f"Procesando párrafo {index}/{total}", index, total)
            file.flush()
            os.fsync(file.fileno())

        _check_worker(worker)
        return ConversionResult.completed(transaction.commit())


def docx_to_pdf(docx_file, output_file, worker=None):
    """Convierte DOCX en PDF simplificado con publicación atómica."""
    output = Path(output_file)
    document = Document(docx_file)
    total = len(document.paragraphs)

    with _single_output_transaction(output) as transaction:
        stage = transaction.stage_for(output)
        pdf_canvas = canvas.Canvas(str(stage), pagesize=A4)
        _, height = A4
        y = height - 50

        for index, paragraph in enumerate(document.paragraphs, start=1):
            _check_worker(worker)
            pdf_canvas.drawString(50, y, paragraph.text)
            y -= 15
            if y < 50:
                pdf_canvas.showPage()
                y = height - 50
            _report(worker, f"Procesando párrafo {index}/{total}", index, total)

        _check_worker(worker)
        pdf_canvas.save()
        _check_worker(worker)
        return ConversionResult.completed(transaction.commit())


# ---------------- TXT Y KML ----------------
def text_to_kml(txt_file, output_file, worker=None):
    """Convierte TXT (lat,lon,nombre opcional) a KML de forma atómica."""
    output = Path(output_file)
    document = XMLDocument()
    kml = document.createElement("kml")
    kml.setAttribute("xmlns", "http://www.opengis.net/kml/2.2")
    document.appendChild(kml)
    root = document.createElement("Document")
    kml.appendChild(root)

    with open(txt_file, "r", encoding="utf-8", errors="ignore") as file:
        for index, line in enumerate(file, start=1):
            _check_worker(worker)
            parts = line.strip().replace("\t", ",").replace(";", ",").split(",")
            if len(parts) >= 2:
                lat, lon = parts[0].strip(), parts[1].strip()
                name = parts[2].strip() if len(parts) > 2 else f"Punto ({lat},{lon})"
                placemark = document.createElement("Placemark")
                pname = document.createElement("name")
                pname.appendChild(document.createTextNode(name))
                placemark.appendChild(pname)
                point = document.createElement("Point")
                coords = document.createElement("coordinates")
                coords.appendChild(document.createTextNode(f"{lon},{lat},0"))
                point.appendChild(coords)
                placemark.appendChild(point)
                root.appendChild(placemark)
            if index % 100 == 0:
                _report(worker, f"Procesando línea {index}")

    _check_worker(worker)
    with _single_output_transaction(output) as transaction:
        stage = transaction.stage_for(output)
        with stage.open("w", encoding="utf-8") as file:
            file.write(document.toprettyxml(indent="  "))
            file.flush()
            os.fsync(file.fileno())
        _check_worker(worker)
        return ConversionResult.completed(transaction.commit())


def kml_to_text(kml_file, output_file, worker=None):
    """Convierte KML a TXT (lat,lon,nombre) mediante staging."""
    from xml.dom import minidom

    output = Path(output_file)
    document = minidom.parse(kml_file)
    placemarks = document.getElementsByTagName("Placemark")
    total = len(placemarks)

    with _single_output_transaction(output) as transaction:
        stage = transaction.stage_for(output)
        with stage.open("w", encoding="utf-8") as file:
            for index, placemark in enumerate(placemarks, start=1):
                _check_worker(worker)
                names = placemark.getElementsByTagName("name")
                name = names[0].firstChild.nodeValue if names else ""
                coords_nodes = placemark.getElementsByTagName("coordinates")
                if not coords_nodes or coords_nodes[0].firstChild is None:
                    raise ValueError(f"Placemark {index} sin coordenadas.")
                coords = coords_nodes[0].firstChild.nodeValue.strip()
                lon, lat, *_ = coords.split(",")
                file.write(f"{lat},{lon},{name}\n")
                _report(worker, f"Procesando punto {index}/{total}", index, total)
            file.flush()
            os.fsync(file.fileno())

        _check_worker(worker)
        return ConversionResult.completed(transaction.commit())


def conversion_task(worker, function, args):
    """Run one conversion and turn ordinary failures into structured results."""
    try:
        worker.check_cancelled()
        result = function(*args, worker=worker)
        if not isinstance(result, ConversionResult):
            result = ConversionResult.completed([str(result)])
        return result
    except WorkerCancelled:
        raise
    except Exception as error:
        logger.exception("Conversion failed")
        return ConversionResult.failed(str(error))


def batch_conversion_task(worker, function, jobs):
    """Convert a folder as all-or-nothing output publication."""
    if not jobs:
        return ConversionResult.failed("No hay archivos para convertir.")

    final_paths = [Path(args[-1]) for args in jobs]
    destination_dir = final_paths[0].parent
    if any(path.parent.resolve(strict=False) != destination_dir.resolve(strict=False) for path in final_paths):
        return ConversionResult.failed("Todos los destinos del lote deben estar en la misma carpeta.")

    warnings = []
    total = len(jobs)
    with OutputTransaction(destination_dir) as transaction:
        try:
            for index, (args, final_path) in enumerate(zip(jobs, final_paths), start=1):
                worker.check_cancelled()
                stage = transaction.stage_for(final_path)
                staged_args = (*args[:-1], str(stage))
                result = function(*staged_args, worker=worker)
                if not isinstance(result, ConversionResult):
                    result = ConversionResult.completed([str(result)])
                warnings.extend(result.warnings)
                if not result.success:
                    failures = result.failures or (f"Falló la conversión de {args[0]}",)
                    return ConversionResult.failed(*failures, warnings=warnings)
                worker.report_progress(
                    {
                        "message": f"Archivo {index}/{total}",
                        "percent": int((index / total) * 100),
                    }
                )

            worker.check_cancelled()
            return ConversionResult.completed(transaction.commit(), warnings=warnings)
        except WorkerCancelled:
            raise
        except Exception as error:
            logger.exception("Batch conversion failed")
            return ConversionResult.failed(str(error), warnings=warnings)


class Tool(BaseTool):
    name = "Convertidor de Archivos"
    description = "Convierte imágenes, PDF, texto, Word y KML."
    category = "Archivos"

    def setup_ui(self):
        self.setWindowTitle(self.name)
        self.setGeometry(200, 200, 420, 460)
        self._worker = None
        self._close_when_worker_finishes = False

        layout = QVBoxLayout()
        buttons = (
            ("Imágenes → PDF", self.convert_images_to_pdf),
            ("PDF → Imágenes", self.convert_pdf_to_images),
            ("TXT → DOCX", self.convert_text_to_docx),
            ("DOCX → TXT", self.convert_docx_to_text),
            ("DOCX → PDF", self.convert_docx_to_pdf),
            ("TXT → KML (archivos o carpeta)", self.convert_text_to_kml),
            ("KML → TXT (archivos o carpeta)", self.convert_kml_to_text),
        )
        for text, callback in buttons:
            button = QPushButton(text)
            button.clicked.connect(callback)
            layout.addWidget(button)

        task_row = QHBoxLayout()
        self.task_status = QLabel("")
        task_row.addWidget(self.task_status)
        task_row.addStretch()
        self.btn_cancel = QPushButton("Cancelar")
        self.btn_cancel.setEnabled(False)
        self.btn_cancel.clicked.connect(self.cancel_conversion)
        task_row.addWidget(self.btn_cancel)
        layout.addLayout(task_row)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

    def _start_conversion(self, label, function, args, success_message):
        if self._worker is not None and self._worker.isRunning():
            QMessageBox.information(self, "Conversión", "Ya hay una conversión en curso.")
            return
        worker = Worker(conversion_task, function, args, parent=self)
        self._start_worker(worker, label, success_message)

    def _start_batch_conversion(self, label, function, jobs, success_message):
        if self._worker is not None and self._worker.isRunning():
            QMessageBox.information(self, "Conversión", "Ya hay una conversión en curso.")
            return
        worker = Worker(batch_conversion_task, function, jobs, parent=self)
        self._start_worker(worker, label, success_message)

    def _start_worker(self, worker, label, success_message):
        self._worker = worker
        self.task_status.setText(label)
        self.btn_cancel.setEnabled(True)
        worker.progress.connect(self._conversion_progress)
        worker.result.connect(lambda result: self._conversion_done(success_message, result))
        worker.error.connect(self._conversion_error)
        worker.cancelled.connect(self._conversion_cancelled)
        worker.finished.connect(lambda: self._conversion_finished(worker))
        worker.start()

    def _conversion_progress(self, progress):
        if isinstance(progress, dict):
            self.task_status.setText(str(progress.get("message", "Procesando...")))
        else:
            self.task_status.setText(str(progress))

    def _conversion_done(self, success_message, result):
        if not isinstance(result, ConversionResult):
            result = ConversionResult.completed([str(result)])

        if result.success:
            message = (
                success_message(list(result.outputs))
                if callable(success_message)
                else success_message
            )
            if result.warnings:
                message += "\n\nAvisos:\n- " + "\n- ".join(result.warnings)
                QMessageBox.warning(self, "Conversión completada con avisos", message)
            else:
                QMessageBox.information(self, "Conversión completada", message)
            return

        message = "No se publicó ningún resultado incompleto."
        if result.failures:
            message += "\n\nErrores:\n- " + "\n- ".join(result.failures)
        if result.warnings:
            message += "\n\nAvisos:\n- " + "\n- ".join(result.warnings)
        QMessageBox.critical(self, "Conversión fallida", message)

    def _conversion_error(self, error):
        logger.error("Conversion failed: %s", error)
        QMessageBox.critical(self, "Error", f"No se pudo completar la conversión:\n{error}")

    def _conversion_cancelled(self):
        self.task_status.setText("Conversión cancelada")

    def _conversion_finished(self, worker):
        if self._worker is worker:
            self._worker = None
            self.btn_cancel.setEnabled(False)
            if self.task_status.text() != "Conversión cancelada":
                self.task_status.setText("")
        worker.deleteLater()
        if self._close_when_worker_finishes:
            self._close_when_worker_finishes = False
            QTimer.singleShot(0, self.close)

    def cancel_conversion(self):
        if self._worker is not None and self._worker.isRunning():
            self._worker.cancel()
            self.btn_cancel.setEnabled(False)
            self.task_status.setText("Cancelando...")

    def closeEvent(self, event):
        if self._worker is not None and self._worker.isRunning():
            self._worker.cancel()
            self._close_when_worker_finishes = True
            self.task_status.setText("Cancelando antes de cerrar...")
            event.ignore()
            return
        super().closeEvent(event)

    def convert_images_to_pdf(self):
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "Seleccionar imágenes",
            "",
            "Imágenes (*.png *.jpg *.jpeg *.bmp)",
        )
        if not files:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Guardar PDF", "", "PDF (*.pdf)")
        if not save_path:
            return
        self._start_conversion(
            "Creando PDF...",
            images_to_pdf,
            (files, save_path),
            f"PDF creado en:\n{save_path}",
        )

    def convert_pdf_to_images(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Seleccionar PDF", "", "PDF (*.pdf)")
        if not file_path:
            return
        folder = QFileDialog.getExistingDirectory(self, "Seleccionar carpeta de destino")
        if not folder:
            return
        self._start_conversion(
            "Renderizando PDF...",
            pdf_to_images,
            (file_path, folder),
            lambda files: f"Se han guardado {len(files)} imágenes en:\n{folder}",
        )

    def convert_text_to_docx(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Seleccionar TXT", "", "Texto (*.txt)")
        if not file_path:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Guardar DOCX", "", "Word (*.docx)")
        if not save_path:
            return
        self._start_conversion(
            "Creando DOCX...",
            text_to_docx,
            (file_path, save_path),
            f"DOCX creado en:\n{save_path}",
        )

    def convert_docx_to_text(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Seleccionar DOCX", "", "Word (*.docx)")
        if not file_path:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Guardar TXT", "", "Texto (*.txt)")
        if not save_path:
            return
        self._start_conversion(
            "Extrayendo texto...",
            docx_to_text,
            (file_path, save_path),
            f"TXT creado en:\n{save_path}",
        )

    def convert_docx_to_pdf(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Seleccionar DOCX", "", "Word (*.docx)")
        if not file_path:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Guardar PDF", "", "PDF (*.pdf)")
        if not save_path:
            return
        self._start_conversion(
            "Creando PDF...",
            docx_to_pdf,
            (file_path, save_path),
            f"PDF creado en:\n{save_path}",
        )

    def convert_text_to_kml(self):
        path = QFileDialog.getExistingDirectory(
            self,
            "Seleccionar carpeta con TXT o archivo individual",
        )
        if path:
            txt_files = [
                os.path.join(path, filename)
                for filename in os.listdir(path)
                if filename.lower().endswith(".txt")
            ]
            if not txt_files:
                QMessageBox.warning(
                    self,
                    "Sin archivos",
                    "No se encontraron archivos .txt en la carpeta seleccionada.",
                )
                return
            output_dir = QFileDialog.getExistingDirectory(
                self,
                "Seleccionar carpeta de destino para KML",
            )
            if not output_dir:
                return
            jobs = []
            for txt_file in txt_files:
                name = os.path.splitext(os.path.basename(txt_file))[0]
                jobs.append((txt_file, os.path.join(output_dir, f"{name}.kml")))
            self._start_batch_conversion(
                "Convirtiendo TXT a KML...",
                text_to_kml,
                jobs,
                lambda outputs: f"Se convirtieron {len(outputs)} archivos a KML.",
            )
            return

        file_path, _ = QFileDialog.getOpenFileName(self, "Seleccionar TXT", "", "Texto (*.txt)")
        if not file_path:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Guardar KML", "", "KML (*.kml)")
        if not save_path:
            return
        self._start_conversion(
            "Convirtiendo TXT a KML...",
            text_to_kml,
            (file_path, save_path),
            f"KML creado en:\n{save_path}",
        )

    def convert_kml_to_text(self):
        path = QFileDialog.getExistingDirectory(
            self,
            "Seleccionar carpeta con KML o archivo individual",
        )
        if path:
            kml_files = [
                os.path.join(path, filename)
                for filename in os.listdir(path)
                if filename.lower().endswith(".kml")
            ]
            if not kml_files:
                QMessageBox.warning(
                    self,
                    "Sin archivos",
                    "No se encontraron archivos .kml en la carpeta seleccionada.",
                )
                return
            output_dir = QFileDialog.getExistingDirectory(
                self,
                "Seleccionar carpeta de destino para TXT",
            )
            if not output_dir:
                return
            jobs = []
            for kml_file in kml_files:
                name = os.path.splitext(os.path.basename(kml_file))[0]
                jobs.append((kml_file, os.path.join(output_dir, f"{name}.txt")))
            self._start_batch_conversion(
                "Convirtiendo KML a TXT...",
                kml_to_text,
                jobs,
                lambda outputs: f"Se convirtieron {len(outputs)} archivos a TXT.",
            )
            return

        file_path, _ = QFileDialog.getOpenFileName(self, "Seleccionar KML", "", "KML (*.kml)")
        if not file_path:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Guardar TXT", "", "Texto (*.txt)")
        if not save_path:
            return
        self._start_conversion(
            "Convirtiendo KML a TXT...",
            kml_to_text,
            (file_path, save_path),
            f"TXT creado en:\n{save_path}",
        )