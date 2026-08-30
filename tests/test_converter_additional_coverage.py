from pathlib import Path

import fitz
import pytest
from docx import Document

from pythonkni.converter import service as converter
from pythonkni.converter.models import ConversionResult
from pythonkni.core.tasks import WorkerCancelled


class RecordingWorker:
    def __init__(self, cancel=False):
        self.cancel = cancel
        self.progress = []

    def check_cancelled(self):
        if self.cancel:
            raise WorkerCancelled()

    def report_progress(self, payload):
        self.progress.append(payload)


def test_output_transaction_rejects_destination_outside_directory(tmp_path):
    transaction = converter.OutputTransaction(tmp_path / "out")
    try:
        with pytest.raises(ValueError, match="compartir carpeta"):
            transaction.stage_for(tmp_path / "other" / "result.txt")
    finally:
        transaction.abort()


def test_output_transaction_missing_stage_fails_and_preserves_destination(tmp_path):
    output = tmp_path / "result.txt"
    output.write_text("previous", encoding="utf-8")
    transaction = converter.OutputTransaction(tmp_path)
    transaction.stage_for(output)

    with pytest.raises(FileNotFoundError, match="resultado temporal"):
        transaction.commit()

    transaction.abort()
    assert output.read_text(encoding="utf-8") == "previous"
    assert not transaction.staging_dir.exists()


def test_output_transaction_success_replaces_existing_and_cleans_staging(tmp_path):
    output = tmp_path / "result.txt"
    output.write_text("old", encoding="utf-8")
    transaction = converter.OutputTransaction(tmp_path)
    stage = transaction.stage_for(output)
    stage.write_text("new", encoding="utf-8")

    assert transaction.commit() == [str(output)]
    assert output.read_text(encoding="utf-8") == "new"
    assert not transaction.staging_dir.exists()

    transaction.abort()
    assert output.read_text(encoding="utf-8") == "new"


def test_worker_helpers_cover_none_and_progress_payloads():
    converter._check_worker(None)
    converter._report(None, "ignored")

    worker = RecordingWorker()
    converter._report(worker, "plain")
    converter._report(worker, "half", current=1, total=2)

    assert worker.progress == [
        {"message": "plain"},
        {"message": "half", "percent": 50},
    ]


def test_pdf_to_images_rejects_empty_pdf_and_closes_document(monkeypatch, tmp_path):
    class EmptyDocument:
        closed = False

        def __len__(self):
            return 0

        def close(self):
            self.closed = True

    document = EmptyDocument()
    monkeypatch.setattr(converter.fitz, "open", lambda _path: document)

    result = converter.pdf_to_images("empty.pdf", tmp_path / "images")

    assert result == ConversionResult.failed("El PDF no contiene páginas para convertir.")
    assert document.closed


def test_text_to_docx_converts_lines_and_reports_progress(tmp_path):
    source = tmp_path / "source.txt"
    output = tmp_path / "result.docx"
    source.write_text("\n".join(f"line {index}" for index in range(101)), encoding="utf-8")
    worker = RecordingWorker()

    result = converter.text_to_docx(source, output, worker=worker)

    assert result.success
    assert result.outputs == (str(output),)
    saved = Document(output)
    assert saved.paragraphs[0].text == "line 0"
    assert saved.paragraphs[-1].text == "line 100"
    assert {"message": "Procesando línea 100"} in worker.progress


def test_docx_to_pdf_creates_output_and_adds_page_when_needed(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "result.pdf"
    document = Document()
    for index in range(60):
        document.add_paragraph(f"paragraph {index}")
    document.save(source)
    worker = RecordingWorker()

    result = converter.docx_to_pdf(source, output, worker=worker)

    assert result.success
    with fitz.open(output) as pdf:
        assert len(pdf) >= 2
    assert worker.progress[-1]["percent"] == 100


def test_text_to_kml_skips_invalid_rows_and_uses_default_name(tmp_path):
    source = tmp_path / "points.txt"
    output = tmp_path / "points.kml"
    source.write_text("38.0,-1.0\ninvalid\n39.0;-2.0;Named", encoding="utf-8")

    result = converter.text_to_kml(source, output)

    assert result.success
    content = output.read_text(encoding="utf-8")
    assert "Punto (38.0,-1.0)" in content
    assert "Named" in content
    assert "-1.0,38.0,0" in content


def test_text_to_kml_reports_each_hundred_lines(tmp_path):
    source = tmp_path / "points.txt"
    output = tmp_path / "points.kml"
    source.write_text("\n".join("38.0,-1.0" for _ in range(100)), encoding="utf-8")
    worker = RecordingWorker()

    converter.text_to_kml(source, output, worker=worker)

    assert {"message": "Procesando línea 100"} in worker.progress


def test_kml_to_text_converts_points(tmp_path):
    source = tmp_path / "points.kml"
    output = tmp_path / "points.txt"
    source.write_text(
        """<?xml version="1.0"?>
<kml><Document>
<Placemark><name>One</name><Point><coordinates>-1.0,38.0,0</coordinates></Point></Placemark>
<Placemark><Point><coordinates>-2.0,39.0,0</coordinates></Point></Placemark>
</Document></kml>""",
        encoding="utf-8",
    )
    worker = RecordingWorker()

    result = converter.kml_to_text(str(source), str(output), worker=worker)

    assert result.success
    assert output.read_text(encoding="utf-8").splitlines() == [
        "38.0,-1.0,One",
        "39.0,-2.0,",
    ]
    assert worker.progress[-1]["percent"] == 100


def test_kml_to_text_missing_coordinates_does_not_replace_existing_output(tmp_path):
    source = tmp_path / "broken.kml"
    output = tmp_path / "points.txt"
    output.write_text("previous", encoding="utf-8")
    source.write_text(
        "<kml><Document><Placemark><name>Broken</name></Placemark></Document></kml>",
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="sin coordenadas"):
        converter.kml_to_text(str(source), str(output))

    assert output.read_text(encoding="utf-8") == "previous"
    assert not list(tmp_path.glob(".pythonkni-converter-*"))


def test_conversion_task_wraps_plain_result(tmp_path):
    worker = RecordingWorker()
    output = tmp_path / "out.txt"

    def convert(source, destination, worker=None):
        assert worker is not None
        Path(destination).write_text(source, encoding="utf-8")
        return destination

    result = converter.conversion_task(worker, convert, ("content", str(output)))

    assert result == ConversionResult.completed([str(output)])


def test_conversion_task_returns_structured_failure():
    worker = RecordingWorker()

    def fail(*_args, **_kwargs):
        raise RuntimeError("boom")

    result = converter.conversion_task(worker, fail, ())

    assert not result.success
    assert result.failures == ("boom",)


def test_conversion_task_propagates_cancellation():
    worker = RecordingWorker(cancel=True)

    with pytest.raises(WorkerCancelled):
        converter.conversion_task(worker, lambda **_kwargs: None, ())


def test_batch_conversion_rejects_empty_and_split_destinations(tmp_path):
    worker = RecordingWorker()
    assert not converter.batch_conversion_task(worker, lambda: None, []).success

    jobs = [
        ("a", str(tmp_path / "one" / "a.txt")),
        ("b", str(tmp_path / "two" / "b.txt")),
    ]
    result = converter.batch_conversion_task(worker, lambda: None, jobs)

    assert not result.success
    assert "misma carpeta" in result.failures[0]


def test_batch_conversion_wraps_plain_results_and_publishes_all(tmp_path):
    worker = RecordingWorker()
    first = tmp_path / "first.txt"
    second = tmp_path / "second.txt"

    def convert(value, destination, worker=None):
        assert worker is not None
        Path(destination).write_text(value, encoding="utf-8")
        return destination

    result = converter.batch_conversion_task(
        worker,
        convert,
        [("one", str(first)), ("two", str(second))],
    )

    assert result.success
    assert first.read_text(encoding="utf-8") == "one"
    assert second.read_text(encoding="utf-8") == "two"
    assert worker.progress[-1] == {"message": "Archivo 2/2", "percent": 100}


def test_batch_conversion_preserves_warnings(tmp_path):
    worker = RecordingWorker()
    output = tmp_path / "result.txt"

    def convert(_value, destination, worker=None):
        Path(destination).write_text("ok", encoding="utf-8")
        return ConversionResult.completed([destination], warnings=["warning"])

    result = converter.batch_conversion_task(worker, convert, [("one", str(output))])

    assert result.success
    assert result.warnings == ("warning",)


def test_batch_conversion_uses_default_failure_message(tmp_path):
    worker = RecordingWorker()
    output = tmp_path / "result.txt"

    def convert(*_args, **_kwargs):
        return ConversionResult(False)

    result = converter.batch_conversion_task(worker, convert, [("source.txt", str(output))])

    assert not result.success
    assert result.failures == ("Falló la conversión de source.txt",)
    assert not output.exists()


def test_batch_conversion_returns_exception_as_failure(tmp_path):
    worker = RecordingWorker()
    output = tmp_path / "result.txt"

    def convert(*_args, **_kwargs):
        raise OSError("conversion exploded")

    result = converter.batch_conversion_task(worker, convert, [("source", str(output))])

    assert not result.success
    assert result.failures == ("conversion exploded",)


def test_batch_conversion_propagates_worker_cancellation(tmp_path):
    worker = RecordingWorker(cancel=True)
    output = tmp_path / "result.txt"

    with pytest.raises(WorkerCancelled):
        converter.batch_conversion_task(
            worker,
            lambda *_args, **_kwargs: ConversionResult.completed([]),
            [("source", str(output))],
        )
