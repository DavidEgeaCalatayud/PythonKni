from pathlib import Path

import fitz
import pytest
from docx import Document
from PIL import Image

from pythonkni.core.tasks import WorkerCancelled
from tools import converter_outputs
from tools.converter_outputs import ConversionResult, OutputTransaction
from tools.converter_tool import (
    batch_conversion_task,
    docx_to_text,
    images_to_pdf,
    pdf_to_images,
    text_to_kml,
    validate_extension,
)


class FakeWorker:
    def __init__(self, cancel_on_check=None):
        self.cancel_on_check = cancel_on_check
        self.check_count = 0

    def check_cancelled(self):
        self.check_count += 1
        if self.cancel_on_check is not None and self.check_count >= self.cancel_on_check:
            raise WorkerCancelled()

    def report_progress(self, _payload):
        return None


def test_validate_extension_accepts_case_insensitive_extensions():
    assert validate_extension("photo.JPG", {".jpg", ".png"})


def test_validate_extension_rejects_unknown_extensions():
    assert not validate_extension("document.exe", {".jpg", ".png"})


def test_images_to_pdf_reports_skipped_images_as_warnings(tmp_path):
    valid_image = tmp_path / "valid.png"
    invalid_image = tmp_path / "invalid.png"
    output = tmp_path / "result.pdf"
    Image.new("RGB", (20, 20), "white").save(valid_image)
    invalid_image.write_text("not an image", encoding="utf-8")

    result = images_to_pdf([str(valid_image), str(invalid_image)], str(output))

    assert isinstance(result, ConversionResult)
    assert result.success is True
    assert result.outputs == (str(output),)
    assert len(result.warnings) == 1
    assert str(invalid_image) in result.warnings[0]
    assert result.failures == ()
    with fitz.open(output) as document:
        assert len(document) == 1


def test_images_to_pdf_does_not_replace_existing_output_when_all_images_fail(tmp_path):
    invalid_image = tmp_path / "invalid.png"
    output = tmp_path / "result.pdf"
    invalid_image.write_text("not an image", encoding="utf-8")
    output.write_bytes(b"previous-valid-output")

    result = images_to_pdf([str(invalid_image)], str(output))

    assert result.success is False
    assert result.outputs == ()
    assert result.failures
    assert result.warnings
    assert output.read_bytes() == b"previous-valid-output"
    assert not list(tmp_path.glob(".pythonkni-converter-*"))


def test_docx_to_text_cancellation_preserves_existing_destination(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "result.txt"
    document = Document()
    document.add_paragraph("first")
    document.add_paragraph("second")
    document.save(source)
    output.write_text("previous", encoding="utf-8")

    worker = FakeWorker(cancel_on_check=2)
    with pytest.raises(WorkerCancelled):
        docx_to_text(str(source), str(output), worker=worker)

    assert output.read_text(encoding="utf-8") == "previous"
    assert not list(tmp_path.glob(".pythonkni-converter-*"))


def test_pdf_to_images_cancellation_rolls_back_entire_page_batch(tmp_path):
    source = tmp_path / "source.pdf"
    output_dir = tmp_path / "images"
    output_dir.mkdir()
    previous_page = output_dir / "page_1.png"
    previous_page.write_bytes(b"previous-page")

    document = fitz.open()
    document.new_page()
    document.new_page()
    document.save(source)
    document.close()

    worker = FakeWorker(cancel_on_check=2)
    with pytest.raises(WorkerCancelled):
        pdf_to_images(str(source), str(output_dir), worker=worker)

    assert previous_page.read_bytes() == b"previous-page"
    assert not (output_dir / "page_2.png").exists()
    assert not list(output_dir.glob(".pythonkni-converter-*"))


def test_batch_conversion_failure_restores_previous_outputs(tmp_path):
    source = tmp_path / "good.txt"
    missing = tmp_path / "missing.txt"
    source.write_text("38.0,-1.0,Good", encoding="utf-8")

    first_output = tmp_path / "good.kml"
    second_output = tmp_path / "missing.kml"
    first_output.write_text("previous-kml", encoding="utf-8")

    result = batch_conversion_task(
        FakeWorker(),
        text_to_kml,
        [
            (str(source), str(first_output)),
            (str(missing), str(second_output)),
        ],
    )

    assert result.success is False
    assert result.failures
    assert first_output.read_text(encoding="utf-8") == "previous-kml"
    assert not second_output.exists()
    assert not list(tmp_path.glob(".pythonkni-converter-*"))


def test_output_transaction_restores_all_destinations_if_publish_fails(monkeypatch, tmp_path):
    first = tmp_path / "first.txt"
    second = tmp_path / "second.txt"
    first.write_text("old-first", encoding="utf-8")
    second.write_text("old-second", encoding="utf-8")

    transaction = OutputTransaction(tmp_path)
    first_stage = transaction.stage_for(first)
    second_stage = transaction.stage_for(second)
    first_stage.write_text("new-first", encoding="utf-8")
    second_stage.write_text("new-second", encoding="utf-8")

    real_replace = converter_outputs.os.replace
    failure_triggered = False

    def flaky_replace(source, destination):
        nonlocal failure_triggered
        if Path(source) == second_stage and not failure_triggered:
            failure_triggered = True
            raise OSError("simulated publish failure")
        return real_replace(source, destination)

    monkeypatch.setattr(converter_outputs.os, "replace", flaky_replace)

    with pytest.raises(OSError, match="simulated publish failure"):
        transaction.commit()
    transaction.abort()

    assert first.read_text(encoding="utf-8") == "old-first"
    assert second.read_text(encoding="utf-8") == "old-second"
    assert not transaction.staging_dir.exists()
